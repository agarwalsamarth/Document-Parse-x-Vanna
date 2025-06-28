import docx
import pandas as pd
import subprocess
import json
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Inches
import vanna as vn
from vanna.remote import VannaDefault
from langchain.llms import Ollama
import sqlite3
import re

vanna_api_key='ff0f0fdd4d5e4ff6ba12a9d3473df087'
vanna_model_name='actuarial_pc_model'
vn_model = VannaDefault(model=vanna_model_name, api_key=vanna_api_key)
vn_model.connect_to_sqlite('/Users/hp/OneDrive/Desktop/Python/SQLITE/Actuarial_PC_Data/Actuarial_PC.db')

def query_vanna(prompt: str) -> tuple[str, pd.DataFrame]:
    sql = vn_model.generate_sql(prompt)
    df = vn_model.run_sql(sql)
    return sql, df

# ----------------------------
# 🤖 Ollama Fuzzy Header Match
# ----------------------------
def get_target_header_and_table(instruction, structure_string):
    prompt = f"""
You are helping identify the correct table to update in a Word document.

The document contains several sections. Each section has a header and a list of tables.
Each table has:
- a table index (starting from 0 under that header),
- the number of rows and columns,
- and the list of column headers (which might vary slightly across documents).

Here is the document structure:
{structure_string}

The user's instruction is:
\"\"\"{instruction}\"\"\"

⚠️ Notes:
- The exact header text may differ slightly from the instruction.
- However, the number and **order** of columns in the correct table will match the instruction's intended schema.
- Choose the table under the best-matching header, with the **most similar column order and count**.

Return your answer strictly in this JSON format:
{{ "header_text": "Exact header from document", "table_index_under_header": 0 }}
"""

    try:
        result = subprocess.run(
            ["ollama", "run", "mistral"],
            input=prompt.encode("utf-8"),
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )
        output = result.stdout.decode("utf-8", errors="replace")
        json_match = re.search(r'\{.*?\}', output, re.DOTALL)
        if json_match:
            return json.loads(json_match.group())
        else:
            print("⚠️ No JSON found in Ollama output:")
            print(output)
            return None
    except Exception as e:
        print("❌ Ollama subprocess failed:", e)
        return None

# ----------------------------
# 📄 Document Parsing
# ----------------------------
def extract_structure(doc):
    structure = []
    current_header = None
    for element in doc.element.body:
        if isinstance(element, CT_P):
            para = Paragraph(element, doc)
            if para.style.name.startswith("Heading"):
                current_header = para.text.strip()
                structure.append({"header": current_header, "tables": []})
        elif isinstance(element, CT_Tbl):
            table = Table(element, doc)
            if current_header is None:
                current_header = "NO_HEADER"
                structure.append({"header": current_header, "tables": []})
            structure[-1]["tables"].append(table)
    return structure

def stringify_structure(structure):
    lines = []
    for section_index, section in enumerate(structure):
        lines.append(f"\n[Section {section_index}] Header: \"{section['header']}\"")
        for i, table in enumerate(section['tables']):
            try:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                lines.append(f"  - Table {i}: {len(table.rows)} rows x {len(table.columns)} cols | Columns: {headers}")
            except Exception:
                lines.append(f"  - Table {i}: Could not extract columns")
    return "\n".join(lines)


# ----------------------------
# 🔁 Replace Table with Formatting
# ----------------------------
def get_column_widths(table: Table):
    widths = []
    for i in range(len(table.columns)):
        try:
            cell = table.cell(0, i)
            widths.append(cell.width)
        except Exception:
            widths.append(None)
    return widths


def replace_table(old_table: Table, new_df: pd.DataFrame):
    parent = old_table._element.getparent()
    index = parent.index(old_table._element)
    widths = get_column_widths(old_table)

    # Remove old table
    parent.remove(old_table._element)

    # Create new table
    doc = old_table._parent
    new_table = doc.add_table(rows=1, cols=len(new_df.columns))
    new_table.style = old_table.style

    # Apply column widths
    for i, col in enumerate(new_table.columns):
        if widths[i]:
            for cell in col.cells:
                cell.width = widths[i]

    # Add header row
    for i, col in enumerate(new_df.columns):
        new_table.cell(0, i).text = str(col)

    # Add data rows
    for _, row in new_df.iterrows():
        cells = new_table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)

    # Insert table at original location
    parent.insert(index, new_table._element)


# ----------------------------
# 🧠 Main Update Loop
# ----------------------------
def main():
    doc_path = '/Users/hp/OneDrive/Desktop/Python/Document Parse x Vanna/Documents/Doc6_With_Two_Tables_One_Header.docx'
    doc = docx.Document(doc_path)

    while True:
        print("\n📝 --- Table Update ---")
        vanna_prompt = input("🔍 Enter Vanna prompt to fetch data: ").strip()
        instruction = input("🛠 Enter update instruction (e.g., 'replace the table under Exposure Year Summary'): ").strip()

        sql, df = query_vanna(vanna_prompt)
        print("\n💡 Vanna SQL:", sql)
        print("📊 Data:\n", df)

        confirm = input("✅ Proceed with this data? (y/n): ").strip().lower()
        if confirm != 'y':
            continue

        structure = extract_structure(doc)
        structure_str = stringify_structure(structure)
        print("\n📘 Document Structure:")
        print(structure_str)

        target = get_target_header_and_table(instruction, structure_str)
        if not target:
            print("⚠️ Could not identify a matching section/table.")
            continue

        header = target['header_text']
        table_idx = target['table_index_under_header']
        print(f"\n🎯 Target: Header='{header}', Table #{table_idx}")

        # Try to find matched section
        matched_section = next((s for s in structure if s['header'] == header), None)

        # If header has no tables, fallback based on column similarity
        if not matched_section or not matched_section['tables']:
            print(f"⚠️ Header '{header}' has no tables. Searching for best matching table...")
            fallback = None
            for section in structure:
                for i, t in enumerate(section['tables']):
                    try:
                        table_headers = [cell.text.strip() for cell in t.rows[0].cells]
                        df_headers = df.columns.tolist()
                        similarity = len(set(table_headers) & set(df_headers)) / max(len(set(df_headers)), 1)
                        if similarity > 0.5:  # threshold can be adjusted
                            fallback = (section, i)
                            break
                    except Exception:
                        continue
                if fallback:
                    break

            if fallback:
                matched_section = fallback[0]
                header = matched_section['header']
                table_idx = fallback[1]
                print(f"➡️ Fallback: Using header '{header}', Table #{table_idx}")
            else:
                print("❌ No similar table found. Skipping update.")
                continue

        # Replace the table
        print(f"📑 Tables under header '{header}': {len(matched_section['tables'])} found.")
        for i, t in enumerate(matched_section['tables']):
            try:
                headers = [cell.text.strip() for cell in t.rows[0].cells]
                print(f"  Table {i}: {len(t.rows)} rows | Columns: {headers}")
            except:
                print(f"  Table {i}: unreadable columns")

        try:
            old_table = matched_section['tables'][table_idx]
            replace_table(old_table, df)
            print("✅ Table replaced successfully.")
        except Exception as e:
            print(f"❌ Error replacing table: {e}")

        again = input("🔁 Update another section? (y/n): ").strip().lower()
        if again != 'y':
            break


    save_path = '/Users/hp/OneDrive/Desktop/Python/Document Parse x Vanna/Documents/Output_Docs2/Doc6_Output3.docx'
    doc.save(save_path)
    print(f"💾 Document saved as: {save_path}")

if __name__ == "__main__":
    main()


#Show me exposure year wise incurred loss, ultimate loss and IBNR
#Update the table under Loss summary by expo year
#Update the exposure year wise incurred, ultimate and ibnr table under tables down here. And also match that the number of columns and headers in the output table should be almost similar with the existing table
#Update the table under Results by Expo year and Profile that includes ExposureYear, TotalIncurredLoss, TotalUltimateLoss, IBNR

#Show me Profile wise avg incurred loss, avg ultimate loss, avg IBNR
#Update the table under losses and ibnr by profile
#Update the profile wise avg incurred, ultimate loss, ibnr table under Tables down here. ALso match the number of columns and the headers in the output table should be almost similar with existing table
#Update the table under Results by Expo year and Profile that has Reserve_Profile, AvgIncurredLoss, AvgUltimateLoss, AvgIBNR

#----------------------------------Flow-----------------------
#✅ Load a Word doc (docx.Document(...))

#🔁 Loop begins

#🧠 Ask the user:

#Vanna prompt → fetches data

#Update instruction → tells where to update

#🔍 Query Vanna with query_vanna(...)

#✅ Show data preview and ask for confirmation

#🗂 Parse document structure (extract_structure)

#🧠 Use Ollama to find matching header/table (get_target_header_and_table)

#🛠 Replace the correct table using replace_table(...)

#🔁 Ask if user wants to update more

#💾 Save the updated doc at the end